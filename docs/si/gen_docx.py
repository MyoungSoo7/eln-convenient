#!/usr/bin/env python3
"""SI 산출물 DOCX 생성 (18건)"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, datetime

OUT = os.path.dirname(os.path.abspath(__file__))
TODAY = datetime.date.today().isoformat()

def shd(cell, color):
    s = cell._element.get_or_add_tcPr()
    s.append(s.makeelement(qn('w:shd'), {qn('w:val'):'clear',qn('w:color'):'auto',qn('w:fill'):color}))

def mk(title, sub=""):
    d = Document()
    for s in ['Normal']+[f'Heading {i}' for i in range(1,4)]:
        st = d.styles[s]; st.font.name = '맑은 고딕'
    d.styles['Normal'].font.size = Pt(10)
    for i in range(1,4): d.styles[f'Heading {i}'].font.color.rgb = RGBColor(0x1E,0x40,0xAF)
    for _ in range(5): d.add_paragraph()
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(title); r.font.size=Pt(26); r.font.bold=True; r.font.color.rgb=RGBColor(0x1E,0x40,0xAF)
    if sub:
        p2=d.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; r2=p2.add_run(sub); r2.font.size=Pt(13); r2.font.color.rgb=RGBColor(0x64,0x74,0x8B)
    d.add_paragraph()
    p3=d.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER; p3.add_run("LabNote ELN 프로젝트").font.size=Pt(13)
    p4=d.add_paragraph(); p4.alignment=WD_ALIGN_PARAGRAPH.CENTER; r4=p4.add_run(f"작성일: {TODAY}  |  버전: 1.0"); r4.font.size=Pt(10); r4.font.color.rgb=RGBColor(0x64,0x74,0x8B)
    d.add_page_break()
    return d

def tbl(d, hdrs, rows):
    t = d.add_table(rows=1+len(rows), cols=len(hdrs)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hdrs):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            for r in p.runs: r.font.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(255,255,255)
        shd(c,'1E40AF')
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
            if ri%2==1: shd(c,'F1F5F9')
    return t

def bul(d, items):
    for i in items: d.add_paragraph(i, style='List Bullet')

def sv(d, name):
    d.save(os.path.join(OUT, name)); print(f"  {name}")

# ════════════════════════════════════════════
# 01. 제안요청서(RFP)
# ════════════════════════════════════════════
def g01():
    d=mk("제안요청서 (RFP)","Request for Proposal")
    d.add_heading("1. 사업 개요",1)
    d.add_paragraph("사내 구축형 전자연구노트(ELN) 협업 플랫폼 구축. 온프레미스 Docker Compose 배포 기반 MSA 아키텍처로, 연구원의 실험 기록 작성, 전자서명, 인벤토리 관리, 장비 예약, 통합 검색 등 연구실 운영 전반을 디지털화한다.")
    d.add_heading("2. 사업 범위",1)
    bul(d,["연구노트 작성/편집/버전관리/실시간 협업","SHA-256 해시체인 기반 전자서명 및 무결성 검증","프로토콜(SOP) 템플릿 관리 및 노트 생성 연동","시약/샘플/장비 인벤토리 관리 (수량, 바코드, 만료 알림)","장비/회의실 예약 시스템 (승인 워크플로)","OpenSearch 기반 통합 검색 (한국어 형태소 분석)","PDF/ZIP 내보내기 (Puppeteer, BullMQ 비동기)","감사 로그 및 실시간 알림","RBAC 기반 권한 관리 (4역할, 18권한)","멀티테넌시 (조직 기반 데이터 격리)","한국어/영어 다국어 지원"])
    d.add_heading("3. 기술 요구사항",1)
    tbl(d,["영역","요구사항"],[["프론트엔드","React + Vite + TypeScript, shadcn/ui, TanStack Query, i18n"],["백엔드","Fastify (전 서비스), Prisma ORM, TypeScript"],["DB","PostgreSQL 15 (단일 인스턴스, 서비스별 스키마 분리)"],["캐시/메시징","Redis 7 (캐시, Stream, Pub/Sub, BullMQ)"],["스토리지","MinIO (S3 호환 오브젝트 스토리지)"],["검색","OpenSearch 2 (한국어 nori 분석기)"],["인증","JWT (jose JWKS + 로컬 듀얼), Keycloak 24 (선택)"],["배포","Docker Compose, 온프레미스"],["모니터링","Jaeger (트레이싱), Dozzle (로그 뷰어)"]])
    d.add_heading("4. 서비스 구성",1)
    tbl(d,["서비스","포트","역할"],[["api-gateway","8000","JWT 검증, 프록시, Rate Limit, SSE, 대시보드"],["auth-service","8001","조직/팀/사용자 CRUD, RBAC, JWT"],["eln-service","8002","노트/프로토콜/템플릿 CRUD, 버전관리"],["signature-audit-service","8003","전자서명, 감사로그, PDF 내보내기, 알림"],["inventory-service","8004","시약/장비 CRUD, 수량관리"],["scheduler-service","8005","예약 CRUD, 승인 워크플로"],["search-service","8006","통합검색, 자동완성"],["file-service","8008","파일 업/다운로드, presigned URL"],["collab-service","8009","WebSocket 실시간 협업"]])
    d.add_heading("5. 납기 및 검수",1)
    d.add_paragraph("납기: 프로젝트 착수일로부터 3개월 이내 최종 검수 완료. 하자보수 기간: 검수 후 12개월.")
    sv(d,"01_제안요청서_RFP.docx")

def g03():
    d=mk("사업수행계획서 (PMP)","Project Management Plan")
    d.add_heading("1. 프로젝트 개요",1)
    tbl(d,["항목","내용"],[["프로젝트명","LabNote ELN (전자연구노트 협업 플랫폼)"],["목표","연구실 실험 기록 디지털화, 전자서명 규정 준수, 협업 효율화"],["범위","8개 MSA 서비스 + 2개 프론트엔드 + 인프라"],["기간","3개월"],["인원","PM 1, DA/TA 1, 개발 6, QA 2, UI 1"]])
    d.add_heading("2. 조직 체계",1)
    tbl(d,["역할","담당","책임"],[["PM","1명","전체 프로젝트 관리, 리스크, 일정, 품질"],["DA/TA","1명","데이터/기술 아키텍처, 설계 검토"],["PL","1명","개발 리드, 코드 리뷰, 기술 의사결정"],["개발","6명","서비스별 개발 (백엔드 4, 프론트 2)"],["QA","2명","테스트 계획, 실행, 결함 관리"],["UI 설계","1명","화면 설계, 프로토타입"]])
    d.add_heading("3. 일정 계획",1)
    tbl(d,["단계","기간","주요 활동","산출물"],[["제안/착수","1주","킥오프, PMP, WBS","제안서, PMP, WBS, 헌장"],["분석","3주","요구사항, 프로세스, 유스케이스","SRS, RTM, BPD"],["설계","3주","아키텍처, ERD, API, 화면","설계서, ERD, API 명세"],["구현","5주","서비스 개발, 단위 테스트","소스코드, 단위테스트 결과"],["테스트","2주","통합/성능/보안/UAT","테스트 결과서"],["이행","1주","배포, 교육, 검수","운영 매뉴얼, 감수 조서"]])
    d.add_heading("4. 위험 관리",1)
    tbl(d,["위험","영향","대응"],[["일정 지연","중","주간 진척 보고, 크리티컬 패스 우선"],["기술 복잡도","중","MSA 경험자 배치, 프로토타입 선행"],["요구사항 변경","중","변경관리 프로세스, RTM 추적"],["인력 이탈","높","문서화, 크로스 트레이닝"]])
    d.add_heading("5. 품질 관리",1)
    bul(d,["코드 리뷰: PR 기반, 최소 1인 승인","단위 테스트: 서비스별 주요 기능 테스트","통합 테스트: 33개 시나리오, 통과율 95% 이상","성능 테스트: 50 동시 사용자, 응답시간 500ms 이내","보안 점검: OWASP Top 10 기반"])
    d.add_heading("6. 커뮤니케이션 계획",1)
    tbl(d,["회의","주기","참석자","목적"],[["스탠드업","매일","전체","진척/이슈 공유"],["주간 보고","주 1회","PM, 발주기관","주간 진척 보고"],["스프린트 리뷰","격주","전체","데모, 피드백"],["월간 보고","월 1회","PM, 경영진","종합 현황 보고"]])
    sv(d,"03_사업수행계획서_PMP.docx")

def g04():
    d=mk("프로젝트 헌장","Project Charter")
    d.add_heading("1. 프로젝트 정보",1)
    tbl(d,["항목","내용"],[["프로젝트명","LabNote ELN (전자연구노트 협업 플랫폼)"],["프로젝트 코드","LABNOTE-ELN-2026"],["스폰서","연구소장"],["PM","프로젝트 매니저"],["시작일",""],["종료일",""],["예산","—"]])
    d.add_heading("2. 프로젝트 목적",1)
    d.add_paragraph("연구실의 종이 기반 실험 기록을 전자화하고, 전자서명을 통한 규정 준수를 보장하며, 시약/장비 관리와 예약 시스템을 통합하여 연구 업무 효율을 극대화한다.")
    d.add_heading("3. 범위",1)
    d.add_heading("3.1 포함",2)
    bul(d,["9개 백엔드 마이크로서비스","2개 프론트엔드 (메인 + 인벤토리 전용)","인프라 구성 (Docker Compose, PostgreSQL, Redis, MinIO, OpenSearch)","사용자 교육 및 운영 매뉴얼"])
    d.add_heading("3.2 제외",2)
    bul(d,["모바일 앱","외부 시스템 연동 (LIMS 등)","하드웨어 조달"])
    d.add_heading("4. 주요 이해관계자",1)
    tbl(d,["이해관계자","역할","관심사"],[["연구소장","스폰서","예산, ROI"],["연구팀장","현업 대표","기능 충족, 사용성"],["연구원","최종 사용자","편의성, 학습 곡선"],["IT 운영팀","시스템 운영","안정성, 유지보수"],["QA팀","품질 보증","결함 관리, 테스트"]])
    d.add_heading("5. 성공 기준",1)
    bul(d,["전체 기능 요구사항 100% 구현","통합 테스트 통과율 95% 이상","성능 목표 달성 (50 동시 사용자, 500ms 응답)","UAT 만족도 4.0/5.0 이상","납기 준수"])
    d.add_heading("6. 승인",1)
    tbl(d,["구분","이름","직책","서명","일자"],[["승인","","스폰서","",""],["확인","","PM","",""]])
    sv(d,"04_프로젝트_헌장.docx")

def g07():
    d=mk("현행시스템 분석서","AS-IS System Analysis")
    d.add_heading("1. 분석 배경",1)
    d.add_paragraph("현재 연구실은 종이 노트, 엑셀, 이메일 등 분산된 도구로 실험 기록을 관리하고 있다. 전자서명 미지원, 검색 불가, 재고 수기 관리 등의 문제가 있어 전자연구노트 시스템 구축이 필요하다.")
    d.add_heading("2. 현행 업무 프로세스",1)
    tbl(d,["업무","현행 방식","문제점"],[["실험 기록","종이 노트 또는 Word 파일","검색 불가, 버전 관리 없음, 훼손 위험"],["서명/검토","수기 서명 또는 날인","위변조 탐지 불가, 추적 어려움"],["시약 관리","엑셀 재고 대장","실시간 파악 불가, 만료 관리 부재"],["장비 예약","구두 또는 공유 캘린더","충돌 빈발, 승인 체계 없음"],["검색","파일 탐색기, 수동 검색","통합 검색 불가, 시간 소요"],["보고서","수동 작성, 인쇄","비효율, 표준화 부재"]])
    d.add_heading("3. 기존 시스템 현황",1)
    d.add_paragraph("전용 ELN 시스템 없음. 일부 연구실에서 범용 도구(MS Office, Google Drive)를 비공식적으로 사용 중. 시약 관리는 연구실별 엑셀 파일로 개별 관리.")
    d.add_heading("4. 개선 방향 (TO-BE)",1)
    tbl(d,["영역","AS-IS","TO-BE"],[["실험 기록","종이/파일","전자 노트, 실시간 협업, 자동 버전관리"],["서명","수기 서명","SHA-256 해시체인 전자서명, 자동 상태 전환"],["시약 관리","엑셀","시스템 CRUD, 바코드, 자동 만료/부족 알림"],["장비 예약","구두","시스템 예약, 승인 워크플로, 충돌 방지"],["검색","수동","OpenSearch 통합검색, 한국어 형태소 분석"],["보고서","수동","자동 PDF/ZIP 생성, 비동기 처리"]])
    sv(d,"07_현행시스템_분석서.docx")

def g11():
    d=mk("유스케이스 명세서","Use Case Specification")
    d.add_heading("1. 액터 정의",1)
    tbl(d,["액터","설명"],[["Researcher","연구원. 노트 작성, 인벤토리 관리, 예약 생성"],["Reviewer","검토자. 노트 검토, 전자서명, 잠금"],["Admin","관리자. 시스템 전체 관리, 잠금 해제, 사용자/역할 관리"],["Viewer","열람자. 읽기 전용"],["System","시스템. 이벤트 기반 자동 처리 (서명→상태전환)"]])
    ucs = [
        ("UC-001","노트 작성","Researcher","1.로그인→2.\"새 노트\" 클릭→3.제목/내용 입력→4.저장→5.상태 draft로 생성","노트가 생성되고 검색 인덱스에 반영","인증 완료, NOTE_WRITE 권한","로그인 미완료 시 /login 리디렉트"),
        ("UC-002","전자서명","Reviewer/Admin","1.\"서명\" 클릭→2.비밀번호 입력→3.해시체인 서명 생성→4.Redis Stream 이벤트→5.노트 signed 전환","노트 상태 signed, 영구 불변","노트 상태 in_progress, note:sign 권한","비밀번호 오류 시 서명 거부"),
        ("UC-003","관리자 잠금해제","Admin","1.잠긴 노트 선택→2.\"잠금 해제\" 클릭→3.비밀번호+사유 입력→4.locked→draft 전환","노트 draft로 복귀, 감사로그+알림 기록","Admin 역할, NOTE_UNLOCK 권한","비밀번호 오류 시 해제 거부"),
        ("UC-004","수량 조정","Researcher","1.아이템 선택→2.수량 조정 버튼→3.유형(in/out/adjust)+수량+사유 입력→4.확인","수량 변경, InventoryHistory 생성","INVENTORY_WRITE 권한","출고 시 재고 부족 에러"),
        ("UC-005","예약 승인","Admin/Owner","1.PENDING 예약 확인→2.\"승인\" 클릭→3.DB 비관적 락→4.충돌 재검증→5.APPROVED 전환→6.알림","예약 승인, 알림 전달","Admin 또는 자원 ownerId","시간 충돌 시 409 에러"),
        ("UC-006","PDF 내보내기","Reviewer/Admin","1.signed 노트 선택→2.\"PDF\" 클릭→3.BullMQ 큐 등록→4.워커 렌더링→5.SSE 알림→6.다운로드","PDF 파일 다운로드 가능","EXPORT_PDF 권한, 노트 signed/locked","렌더링 실패 시 3회 재시도"),
    ]
    for uc in ucs:
        d.add_heading(f"{uc[0]}: {uc[1]}",2)
        tbl(d,["항목","내용"],[["유스케이스 ID",uc[0]],["유스케이스명",uc[1]],["액터",uc[2]],["주요 흐름",uc[3]],["사후 조건",uc[4]],["사전 조건",uc[5]],["예외 흐름",uc[6]]])
        d.add_paragraph()
    sv(d,"11_유스케이스_명세서.docx")

def g12():
    d=mk("인터뷰/회의록","Meeting Minutes")
    d.add_heading("1. 킥오프 미팅",1)
    tbl(d,["항목","내용"],[["일시",""],["장소","회의실"],["참석자","PM, DA, PL, 개발팀, QA, UI, 발주기관"],["안건","프로젝트 착수, 범위 확인, 일정 합의"]])
    d.add_heading("주요 논의 사항",2)
    bul(d,["전자서명 방식: SHA-256 해시체인 채택 (PKI 대비 구현 용이, 온프레미스 적합)","MSA 서비스 분리 기준: 비즈니스 도메인 단위 (auth, eln, signature, inventory, scheduler, search, file, collab)","DB 전략: 단일 PostgreSQL, 서비스별 Prisma 스키마 분리 (운영 단순화)","검색 엔진: OpenSearch + 한국어 nori 분석기 채택","실시간 협업: WebSocket + Redis Pub/Sub 팬아웃","PDF 내보내기: Puppeteer + BullMQ 비동기 처리 (서비스 블로킹 방지)","다국어: react-i18next (한국어/영어)"])
    d.add_heading("결정 사항",2)
    bul(d,["Fastify로 전 서비스 통일 (Express 대비 성능 우위)","Keycloak SSO는 선택 기능으로 (없이도 동작)","멀티테넌시: orgId 기반 데이터 격리 (모든 쿼리 스코핑)","Docker Compose 온프레미스 배포 (K8s 미사용, 운영 단순화)"])
    d.add_heading("2. 요구사항 확인 회의",1)
    tbl(d,["항목","내용"],[["일시",""],["참석자","PM, AA, 연구팀장, 연구원 대표"],["안건","기능 요구사항 확인, 우선순위 합의"]])
    d.add_heading("주요 논의 사항",2)
    bul(d,["노트 상태 흐름 확정: draft↔in_progress→locked/signed, signed 불변","Researcher는 서명 불가 (규정 준수: 작성자와 서명자 분리)","인벤토리 유형은 코드 테이블로 동적 관리 (하드코딩 X)","예약 승인 시 비관적 락 필요 (동시 승인 방지)","대시보드: 개인/팀/조직 3레벨 제공"])
    sv(d,"12_인터뷰_회의록.docx")

def g13():
    d=mk("분석단계 완료보고서","Analysis Phase Completion Report")
    d.add_heading("1. 분석 결과 요약",1)
    tbl(d,["항목","내용"],[["분석 기간","2주"],["주요 산출물","현행시스템 분석서, SRS, RTM, BPD, 유스케이스 명세서"],["기능 요구사항","28건 (FR-001 ~ FR-028)"],["비기능 요구사항","12건 (NFR-001 ~ NFR-012)"],["업무 프로세스","8건 (BP-001 ~ BP-008)"],["유스케이스","6건 (UC-001 ~ UC-006)"]])
    d.add_heading("2. 주요 분석 결과",1)
    bul(d,["9개 서비스 도메인 식별 (auth, eln, sig-audit, inventory, scheduler, search, file, collab, gateway)","핵심 비즈니스 규칙 정의: 노트 상태 전환 맵, 서명 프로세스, 예약 충돌 검증","서비스 간 통신 패턴 정의: Redis Stream (비동기), HTTP (동기), Pub/Sub (실시간)","멀티테넌시 전략 확정: orgId 기반 전 서비스 일관 적용"])
    d.add_heading("3. 이슈 및 조치",1)
    tbl(d,["이슈","영향","조치"],[["서명 방식 선택","높음","SHA-256 해시체인 채택 (PKI 복잡도 회피)"],["검색 엔진 선택","중","OpenSearch 채택 (ElasticSearch 라이선스 이슈 회피)"],["실시간 협업 범위","중","텍스트 기반 협업 우선 (CRDT는 v2로 연기)"]])
    d.add_heading("4. 설계 단계 이행 사항",1)
    bul(d,["요구사항 정의서 기반 시스템 아키텍처 설계 착수","데이터 모델(ERD) 설계: 7개 스키마 분리","API 명세서 작성: 서비스별 OpenAPI 스펙","화면 설계: 15개 페이지 와이어프레임"])
    sv(d,"13_분석단계_완료보고서.docx")

def g14():
    d=mk("시스템 아키텍처 설계서","System Architecture Design")
    d.add_heading("1. 아키텍처 개요",1)
    d.add_paragraph("LabNote ELN은 MSA(Microservice Architecture) 기반으로 설계되었다. 9개 백엔드 서비스가 Docker Compose로 배포되며, API Gateway를 통해 외부 요청을 라우팅한다.")
    d.add_heading("2. 서비스 아키텍처",1)
    tbl(d,["서비스","포트","역할","주요 기술","DB 스키마"],[
        ["api-gateway","8000","JWT 검증, 프록시, Rate Limit, SSE, 대시보드 집계","Fastify, @fastify/http-proxy, ioredis","—"],
        ["auth-service","8001","조직/팀/사용자 CRUD, RBAC, JWT 발급","Fastify, Prisma, jose, bcrypt","auth"],
        ["eln-service","8002","노트/프로토콜/템플릿 CRUD, 버전관리, 이벤트 소비","Fastify, Prisma, ioredis","eln"],
        ["signature-audit","8003","전자서명(해시체인), 감사로그, PDF 내보내기, 알림","Fastify, Prisma, BullMQ, Puppeteer","signature"],
        ["inventory-service","8004","시약/장비 CRUD, 수량관리, 바코드","Fastify, Prisma","inventory"],
        ["scheduler-service","8005","예약 CRUD, 승인/거절, 비관적 락","Fastify, Prisma","scheduler"],
        ["search-service","8006","통합검색, 자동완성, 인덱스 관리","Fastify, Prisma, @opensearch-project/opensearch","search"],
        ["file-service","8008","파일 업/다운로드, presigned URL","Fastify, Prisma, @aws-sdk/client-s3","file"],
        ["collab-service","8009","WebSocket 실시간 협업, Redis Pub/Sub","ws, ioredis, jsonwebtoken","—"],
    ])
    d.add_heading("3. 통신 패턴",1)
    tbl(d,["패턴","기술","용도","예시"],[
        ["동기 HTTP","fetch/http.request + x-internal-secret","서비스 간 직접 호출","비밀번호 검증, 감사로그 기록, 검색 인덱스"],
        ["비동기 이벤트","Redis Stream (XADD/XREADGROUP)","이벤트 기반 상태 전환","NOTE_SIGNED: sig-audit→eln"],
        ["실시간 알림","Redis Pub/Sub → SSE","사용자 실시간 알림","export-status, notifications:{userId}"],
        ["백그라운드 작업","BullMQ (Redis 기반)","비동기 작업 큐","PDF/ZIP 내보내기"],
        ["실시간 협업","WebSocket + Redis Pub/Sub","다중 사용자 동시 편집","collab-service 팬아웃"],
        ["API 프록시","@fastify/http-proxy","외부→내부 라우팅","api-gateway → 각 서비스"],
    ])
    d.add_heading("4. 인증/인가 아키텍처",1)
    d.add_paragraph("1) api-gateway가 JWT 검증 (jose JWKS + 로컬 듀얼)\n2) 검증 후 헤더 주입: x-user-id, x-user-role, x-user-permissions, x-user-org-id\n3) 각 서비스: requireAuth → requirePermission 미들웨어\n4) 내부 통신: x-internal-secret 헤더 인증\n5) 모든 쿼리: getOrgId() + withOrgScope() 멀티테넌시 격리")
    d.add_heading("5. 인프라 아키텍처",1)
    tbl(d,["컴포넌트","버전","용도","설정"],[
        ["PostgreSQL","15","관계형 DB","단일 인스턴스, 7개 스키마 분리"],
        ["Redis","7","캐시, Stream, Pub/Sub, BullMQ","단일 인스턴스"],
        ["MinIO","latest","오브젝트 스토리지","S3 호환, labnote-* 버킷"],
        ["OpenSearch","2","검색 엔진","한국어 nori 분석기, labnote_notes 인덱스"],
        ["Keycloak","24","SSO (선택)","realm-labnote.json"],
        ["Jaeger","latest","분산 트레이싱","OpenTelemetry"],
        ["Dozzle","latest","로그 뷰어","Docker 로그 집계"],
    ])
    sv(d,"14_시스템_아키텍처_설계서.docx")

def g17():
    d=mk("인터페이스 정의서","Interface Specification")
    d.add_heading("1. 내부 서비스 간 인터페이스",1)
    tbl(d,["호출자","대상","Method","경로","인증","용도","데이터"],[
        ["sig-audit","auth","POST","/api/auth/internal/verify-password","x-internal-secret","서명 전 비밀번호 검증","{userId,password}→{verified}"],
        ["sig-audit","eln","GET","/api/notes/:id","x-internal-secret","서명 시 노트 내용 조회","Note 전문"],
        ["sig-audit","eln","PATCH","/api/notes/:id/status","x-internal-secret","Redis 실패 시 HTTP 폴백","{status:'signed'}"],
        ["sig-audit","file","POST","/api/exports/internal/upload","x-internal-secret","PDF/ZIP MinIO 업로드","multipart"],
        ["eln","auth","POST","/api/auth/internal/verify-password","x-internal-secret","관리자 잠금해제 비밀번호 확인","{userId,password}"],
        ["eln","sig-audit","POST","/api/audit/internal","x-internal-secret","감사 로그 기록","{entityType,action,actorId...}"],
        ["eln","sig-audit","POST","/api/notifications/internal","x-internal-secret","잠금/해제 알림 발송","{recipientId,type,title...}"],
        ["eln","search","POST","/api/search/index","x-internal-secret","노트/템플릿 검색 인덱스 갱신","{id,doc}"],
        ["auth","sig-audit","POST","/api/audit/internal","x-internal-secret","인증 관련 감사 로그","{action:login/logout...}"],
        ["inventory","search","POST","/api/search/index","x-internal-secret","인벤토리 검색 인덱스 갱신","{id,doc}"],
        ["scheduler","sig-audit","POST","/api/notifications/internal","x-internal-secret","예약 승인 알림","{type:BOOKING_APPROVED}"],
    ])
    d.add_heading("2. 이벤트 기반 인터페이스",1)
    tbl(d,["이벤트","발행자","구독자","채널/Stream","데이터","비고"],[
        ["NOTE_SIGNED","sig-audit","eln","Redis Stream: labnote:events","{noteId,status,userId}","소비자 그룹: eln-service"],
        ["export-status","sig-audit worker","gateway SSE","Redis Pub/Sub: export-status","{jobId,status,fileUrl}","사용자별 필터링"],
        ["notification","sig-audit","gateway SSE","Redis Pub/Sub: notifications:{userId}","{id,type,title,message}","사용자별 채널"],
        ["collab","collab-service","collab-service","Redis Pub/Sub: labnote:collab","{noteId,payload,sourceUserId}","다중 인스턴스 팬아웃"],
    ])
    d.add_heading("3. 외부 인터페이스",1)
    d.add_paragraph("현재 외부 시스템 연동 없음. Keycloak SSO는 선택적 연동 (JWKS 엔드포인트).")
    sv(d,"17_인터페이스_정의서.docx")

def g27():
    d=mk("구현단계 완료보고서","Implementation Phase Report")
    d.add_heading("1. 구현 결과 요약",1)
    tbl(d,["항목","내용"],[["구현 기간","5주"],["서비스 수","9개 백엔드 + 2개 프론트엔드"],["프로그램 수","22개 모듈 (PG-001 ~ PG-022)"],["소스코드","Git Repository"],["단위 테스트","32건 전체 PASS"],["코드 인스펙션","16개 항목 점검, 3건 개선"]])
    d.add_heading("2. 서비스별 구현 현황",1)
    tbl(d,["서비스","주요 기능","API 수","상태"],[
        ["auth-service","로그인/JWT/RBAC/조직/팀/사용자/역할 관리","26","완료"],
        ["eln-service","노트/프로토콜/템플릿/코드 CRUD, 상태 전환, 이벤트 소비","34","완료"],
        ["signature-audit","전자서명, 감사로그, PDF 내보내기, 알림","18","완료"],
        ["inventory-service","아이템 CRUD, 수량 관리, 만료/재고 알림","14","완료"],
        ["scheduler-service","자원/예약 CRUD, 승인 워크플로","15","완료"],
        ["search-service","통합검색, 자동완성, 즐겨찾기, 인덱스","14","완료"],
        ["file-service","파일 업/다운로드, presigned URL, 내보내기","14","완료"],
        ["collab-service","WebSocket 실시간 협업","—","완료"],
        ["api-gateway","JWT 프록시, Rate Limit, SSE, 대시보드","9","완료"],
        ["프론트엔드","15개 페이지, i18n, 다크모드","—","완료"],
    ])
    d.add_heading("3. 발견 이슈 및 조치",1)
    tbl(d,["이슈","조치"],[["비밀번호 정책 미적용 (DEF-006)","Zod 스키마에 정책 추가 완료"],["검색 인덱스 삭제 누락 (DEF-001)","searchClient.delete() 호출 추가 완료"],["XSS 위험 (검색 하이라이트)","DOMPurify 적용 완료"]])
    sv(d,"27_구현단계_완료보고서.docx")

def g37():
    d=mk("운영 이행 계획서","Operations Transition Plan")
    d.add_heading("1. 이행 개요",1)
    d.add_paragraph("LabNote ELN 시스템을 운영 환경에 배포하기 위한 계획. Docker Compose 기반 온프레미스 배포.")
    d.add_heading("2. 이행 절차",1)
    tbl(d,["단계","작업","담당","확인사항"],[
        ["1","운영 서버 환경 준비 (Docker, 네트워크)","DevOps","Docker Engine 24+, Compose v2"],
        ["2",".env 파일 구성 (DB, Redis, MinIO, JWT 시크릿)","DevOps","시크릿 값 생성, 환경변수 검증"],
        ["3","docker compose up -d (전체 서비스 기동)","DevOps","전 서비스 healthy 상태 확인"],
        ["4","Prisma migrate deploy (DB 스키마 적용)","DevOps","마이그레이션 성공 확인"],
        ["5","Prisma db seed (초기 데이터 투입)","DevOps","관리자 계정, 기본 역할/권한"],
        ["6","OpenSearch 인덱스 생성 및 매핑","DevOps","nori 분석기, 매핑 확인"],
        ["7","MinIO 버킷 생성","DevOps","labnote-files, labnote-exports"],
        ["8","기존 데이터 마이그레이션 (해당 시)","DA","데이터 정합성 확인"],
        ["9","연동 테스트 (헬스체크, 기본 기능)","QA","전체 /health 정상"],
        ["10","사용자 교육","PM","교육 완료 확인"],
    ])
    d.add_heading("3. 롤백 계획",1)
    bul(d,["Docker 이미지 태그 관리: 이전 버전 이미지 보존","DB 백업: 이행 전 pg_dump 실행","롤백 절차: docker compose down → 이전 이미지 태그로 재기동 → DB 복원"])
    d.add_heading("4. 운영 환경 정보",1)
    tbl(d,["항목","사양"],[["OS","Ubuntu 22.04 LTS"],["CPU","8 vCPU"],["메모리","16GB RAM"],["디스크","SSD 200GB"],["Docker","Engine 24+, Compose v2"],["네트워크","내부 네트워크 (외부 접근: 5173, 8000)"]])
    sv(d,"37_운영_이행_계획서.docx")

def g38():
    d=mk("운영자 매뉴얼","Operations Manual")
    d.add_heading("1. 시스템 구성",1)
    d.add_paragraph("LabNote ELN은 9개 백엔드 서비스 + 인프라 4종으로 구성된 Docker Compose 환경이다.")
    d.add_heading("2. 서비스 관리 명령어",1)
    tbl(d,["명령어","설명"],[["docker compose up -d","전체 서비스 시작"],["docker compose up -d --build <서비스>","특정 서비스 재빌드 후 시작"],["docker compose down","전체 서비스 중지"],["docker compose logs -f <서비스>","로그 확인"],["docker compose ps","서비스 상태 확인"],["docker compose restart <서비스>","서비스 재시작"]])
    d.add_heading("3. 헬스체크",1)
    d.add_paragraph("각 서비스는 /health 엔드포인트를 제공한다. Docker healthcheck로 자동 모니터링.")
    tbl(d,["서비스","URL"],[["api-gateway","http://localhost:8000/health"],["auth-service","http://localhost:8001/health"],["eln-service","http://localhost:8002/health"],["signature-audit","http://localhost:8003/health"],["inventory-service","http://localhost:8004/health"],["scheduler-service","http://localhost:8005/health"],["search-service","http://localhost:8006/health"],["file-service","http://localhost:8008/health"]])
    d.add_heading("4. DB 관리",1)
    tbl(d,["작업","명령어"],[["마이그레이션","docker exec <컨테이너> npx prisma migrate deploy"],["시드 데이터","docker exec <컨테이너> npx prisma db seed"],["DB GUI","docker exec <컨테이너> npx prisma studio"],["백업","docker exec postgres pg_dump -U labnote > backup.sql"],["복원","docker exec -i postgres psql -U labnote < backup.sql"]])
    d.add_heading("5. 모니터링",1)
    tbl(d,["도구","URL","용도"],[["Jaeger","http://localhost:16686","분산 트레이싱"],["Dozzle","http://localhost:9999","로그 뷰어"],["MinIO 콘솔","http://localhost:9001","파일 스토리지 관리"]])
    d.add_heading("6. 장애 대응",1)
    tbl(d,["증상","원인","대응"],[["서비스 unhealthy","DB/Redis 연결 실패","docker compose restart <서비스>, 인프라 상태 확인"],["PDF 내보내기 실패","Puppeteer/BullMQ 이슈","sig-audit 로그 확인, 워커 재시작"],["검색 결과 없음","OpenSearch 인덱스 손상","인덱스 재생성 (search-service /api/search/index/bulk)"],["Redis Stream 적체","소비자 지연","eln-service 재시작, pending 메시지 확인"]])
    sv(d,"38_운영자_매뉴얼.docx")

def g39():
    d=mk("사용자 매뉴얼","User Manual")
    d.add_heading("1. 시스템 접속",1)
    d.add_paragraph("브라우저에서 http://localhost:5173 접속 → 이메일/비밀번호 입력 → 로그인")
    d.add_heading("2. 주요 기능 안내",1)
    tbl(d,["메뉴","경로","설명"],[["대시보드","/","통계, 최근 노트, 예약, 알림"],["연구노트","/notes","노트 CRUD, 상태 전환, 첨부/연결"],["프로토콜","/protocols","템플릿 관리, 노트 생성 연동"],["인벤토리","/inventory","시약/장비 관리, 수량 조정"],["예약","/scheduler","장비/회의실 예약, 캘린더"],["검색","/search","통합 검색, 즐겨찾기"],["서명 현황","/signatures","전자서명 관리"],["감사 로그","/audit-logs","활동 기록 조회"],["내보내기","/exports","PDF/ZIP 내보내기"]])
    d.add_heading("3. 노트 상태 흐름",1)
    d.add_paragraph("draft(초안) ↔ in_progress(진행중) → locked(잠김) 또는 signed(서명완료)\n- signed 상태는 영구 불변 (수정/삭제/상태변경 불가)\n- locked→draft는 Admin만 가능 (비밀번호+사유 필요)")
    d.add_heading("4. 역할별 권한",1)
    tbl(d,["기능","Admin","Reviewer","Researcher","Viewer"],[
        ["노트 작성/수정","O","O","O","X"],["전자서명","O","O","X","X"],
        ["잠금 해제","O","X","X","X"],["관리자 메뉴","O","X","X","X"],
        ["예약 승인","O","O","X","X"]])
    d.add_paragraph("(상세 사용법은 별도 사용자 매뉴얼 PPT 참고)")
    sv(d,"39_사용자_매뉴얼.docx")

def g41():
    d=mk("하자보수 계획서","Warranty & Maintenance Plan")
    d.add_heading("1. 하자보수 개요",1)
    tbl(d,["항목","내용"],[["하자보수 기간","최종 검수일로부터 12개월"],["대상 시스템","LabNote ELN 전체 (9개 서비스 + 인프라)"],["담당","개발팀 + DevOps"]])
    d.add_heading("2. SLA 정의",1)
    tbl(d,["등급","정의","응답 시간","해결 시간"],[
        ["Critical","시스템 전체 장애, 데이터 유실 위험","1시간 이내","4시간 이내"],
        ["High","주요 기능 장애 (서명, 노트 저장 등)","2시간 이내","8시간 이내"],
        ["Medium","부가 기능 장애 (검색, 알림 등)","4시간 이내","24시간 이내"],
        ["Low","UI 오류, 개선 요청","1영업일","5영업일"]])
    d.add_heading("3. 지원 범위",1)
    d.add_heading("3.1 포함",2)
    bul(d,["소프트웨어 결함 수정","보안 패치 적용","Docker 이미지 업데이트","DB 마이그레이션 지원","장애 원인 분석 및 조치","운영 문의 대응"])
    d.add_heading("3.2 제외",2)
    bul(d,["신규 기능 개발 (별도 계약)","하드웨어 장애","사용자 과실에 의한 데이터 손실","네트워크/인프라 장애 (고객 책임)"])
    d.add_heading("4. 연락 채널",1)
    tbl(d,["채널","용도"],[["이메일","일반 문의, 결함 신고"],["전화","Critical/High 긴급 장애"],["이슈 트래커","결함 등록, 진행 추적"]])
    sv(d,"41_하자보수_계획서.docx")

def g43():
    d=mk("프로젝트 완료보고서","Project Completion Report")
    d.add_heading("1. 프로젝트 요약",1)
    tbl(d,["항목","내용"],[["프로젝트명","LabNote ELN (전자연구노트 협업 플랫폼)"],["목표","연구실 실험 기록 디지털화, 전자서명 규정 준수"],["기간","3개월"],["팀 규모","PM 1, DA/TA 1, 개발 6, QA 2, UI 1"],["산출물","44건 (최종 산출물 목록 참조)"]])
    d.add_heading("2. 목표 달성 현황",1)
    tbl(d,["목표","달성","비고"],[
        ["기능 요구사항 100% 구현","O","28개 FR 전체 구현"],
        ["통합 테스트 통과율 95%+","O","33개 시나리오 100% PASS"],
        ["성능 목표 달성","O (조건부)","PDF 내보내기 28.5초 (목표 30초 근접)"],
        ["보안 점검 통과","O","Critical 0건, Medium 미조치 1건 (비밀번호 정책→완료)"],
        ["UAT 만족도 4.0+","O","평균 4.2/5.0"],
        ["납기 준수","O",""]])
    d.add_heading("3. 산출물 현황",1)
    tbl(d,["단계","산출물 수","완료","비고"],[["제안","6","6",""],["분석","7","7",""],["설계","8+","8","화면설계서 별도"],["구현","6","6","소스코드 Git"],["테스트","9","9",""],["이행/완료","8","8",""],["합계","44+","44",""]])
    d.add_heading("4. 결함 현황",1)
    tbl(d,["심각도","발견","종결","미종결"],[["Critical","0","0","0"],["Medium","3","3","0"],["Low","5","4","1 (보류: Enhancement)"],["합계","9","7","2 (1 진행중, 1 보류)"]])
    d.add_heading("5. 인수인계 사항",1)
    bul(d,["운영자 매뉴얼 및 사용자 매뉴얼 인계 완료","운영 환경 접속 정보 및 시크릿 인계","하자보수 계획서 합의 (12개월, SLA 포함)","교육 완료 (연구원 3명, 검토자 1명, 관리자 1명)"])
    sv(d,"43_프로젝트_완료보고서.docx")

def g44():
    d=mk("프로젝트 감수 조서","Acceptance Certificate")
    d.add_heading("1. 감수 대상",1)
    tbl(d,["항목","내용"],[["프로젝트명","LabNote ELN (전자연구노트 협업 플랫폼)"],["납품 범위","9개 백엔드 서비스, 2개 프론트엔드, 인프라, 산출물 44건"],["감수 기간",""],["감수 장소",""]])
    d.add_heading("2. 감수 항목 및 결과",1)
    tbl(d,["No","감수 항목","기준","결과","비고"],[
        ["1","산출물 완전성","44건 산출물 목록 대비 누락 없음","적합",""],
        ["2","기능 요구사항","28건 FR 전체 구현 및 동작 확인","적합",""],
        ["3","비기능 요구사항","12건 NFR 충족 확인","적합","PDF 내보내기 조건부"],
        ["4","통합 테스트","33개 시나리오 통과율 95%+","적합","100% PASS"],
        ["5","성능 테스트","50 동시 사용자, 응답시간 500ms 이내","적합",""],
        ["6","보안 점검","OWASP Top 10 기준 Critical/High 0건","적합",""],
        ["7","UAT","만족도 4.0/5.0 이상","적합","평균 4.2"],
        ["8","운영 환경 배포","정상 기동 및 헬스체크 확인","적합",""],
        ["9","교육 완료","사용자/운영자 교육 실시","적합",""],
        ["10","매뉴얼","운영자/사용자 매뉴얼 제공","적합",""]])
    d.add_heading("3. 종합 판정",1)
    p=d.add_paragraph(); r=p.add_run("판정: 적합 (검수 통과)"); r.font.size=Pt(14); r.font.bold=True; r.font.color.rgb=RGBColor(0x10,0xB9,0x81)
    d.add_heading("4. 조건 및 유의사항",1)
    bul(d,["PDF 내보내기 성능 개선 (Puppeteer 최적화) — 하자보수 기간 내 조치","HTTP 보안 헤더 적용 (DEF-007) — 하자보수 기간 내 조치","향후 기능 개선 요청 (캘린더 클릭 예약, 일괄 입고) — 별도 계약"])
    d.add_heading("5. 서명",1)
    tbl(d,["구분","이름","직책","서명","일자"],[["납품자","","PM","",""],["검수자","","발주기관 대표","",""],["입회자","","QA 리드","",""]])
    sv(d,"44_프로젝트_감수_조서.docx")

# ── 실행 ──
print("=== DOCX 생성 시작 ===")
g01(); g03(); g04(); g07(); g11(); g12(); g13(); g14(); g17(); g27(); g37(); g38(); g39(); g41(); g43(); g44()
print("=== DOCX 16건 완료 ===")
